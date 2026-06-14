#!/usr/bin/env python3
"""Generate line-numbered source-only DOCX for CIT paper (current version)."""
import os, sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

BASE = '/Users/caihengjin/.openclaw/workspace/analysis'

def create_source_only(input_tex_path, output_name, title):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(7)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(8.5)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(6)

    # Read and number lines
    with open(input_tex_path, 'r') as f:
        content = f.read()

    lines = content.split('\n')
    numbered = [f"{i+1}|{line}" for i, line in enumerate(lines)]

    # Write in batches
    batch_size = 50
    for i in range(0, len(numbered), batch_size):
        batch = '\n'.join(numbered[i:i+batch_size])
        p = doc.add_paragraph()
        run = p.add_run(batch)
        run.font.name = 'Courier New'
        run.font.size = Pt(7)

    path = os.path.join(BASE, output_name)
    doc.save(path)
    print(f"Saved: {path} ({len(lines)} lines)")

print("Generating CIT paper DOCX from current sources...")

# Read preamble and body
with open(os.path.join(BASE, 'preamble.tex'), 'r') as f:
    preamble = f.read()

with open(os.path.join(BASE, 'CIT_body.tex'), 'r') as f:
    body = f.read()

# Combine into full document
full_tex = preamble + '\n\\begin{document}\n' + body + '\n\\bibliography{rat_refs}\n\\end{document}\n'

# Write temp file
temp_path = os.path.join(BASE, '_cit_full.tex')
with open(temp_path, 'w') as f:
    f.write(full_tex)

# Create main text DOCX
create_source_only(temp_path, 'CIT_paper_source_only.docx',
                   'Main Text — Cognitive Inertia Theorem')

# Create supplementary DOCX
create_source_only(os.path.join(BASE, 'CIT_supplementary.tex'),
                   'CIT_supplementary_source_only.docx',
                   'Supplementary Material — Cognitive Inertia Theorem')

os.remove(temp_path)
print("Done!")
